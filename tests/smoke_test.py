from __future__ import annotations

import shutil
import sys
from pathlib import Path

from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from watermark_app.docx_adapter import embed_docx_keep, extract_docx_keep
from watermark_app.image_watermark import embed_image, extract_image


TMP = ROOT / "tmp" / "smoke"


def make_image(path: Path) -> None:
    image = Image.new("RGB", (960, 720), "white")
    draw = ImageDraw.Draw(image)
    for y in range(0, 720, 24):
        draw.line((0, y, 960, y), fill=(220, 225, 232))
    for x in range(0, 960, 24):
        draw.line((x, 0, x, 720), fill=(235, 238, 242))
    draw.rectangle((120, 120, 840, 600), outline=(70, 100, 140), width=4)
    draw.text((160, 180), "Invisible watermark smoke image", fill=(40, 50, 70))
    image.save(path)


def make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Smoke Test Document", 0)
    doc.add_paragraph("This document is used to test retained DOCX watermark payloads.")
    doc.save(path)


def main() -> int:
    if TMP.exists():
        shutil.rmtree(TMP)
    TMP.mkdir(parents=True)
    password = "test-password"
    text = "内部资料，仅授权张三使用。项目编号 WM-2026-001。未经许可不得外传。"

    source_image = TMP / "source.png"
    make_image(source_image)
    embedded = embed_image(source_image, TMP, text, password)
    recovered = extract_image(embedded.output_path, password)
    assert recovered.payload.core_text == ""
    assert recovered.payload.watermark_id == embedded.watermark_id
    assert recovered.algorithm == "image-iwm2-local"
    assert recovered.tiles_verified >= 1

    source_docx = TMP / "source.docx"
    make_docx(source_docx)
    docx_result = embed_docx_keep(source_docx, TMP, text, password)
    docx_payload = extract_docx_keep(docx_result.output_path, password)
    assert docx_payload.core_text == text

    print("Smoke tests passed.")
    print(f"Image: {embedded.output_path}")
    print(f"DOCX: {docx_result.output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
